#coding=utf-8

from .models import QuestionBank
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import platform
from django.conf import settings
import os
import random
from urllib import request
from io import StringIO, BytesIO
import urllib
from PIL import Image


def exportword(choose):
    # 打开文档
    document = Document()

    # 加入不同等级的标题
    document.add_heading('多益网络IT工程师笔试题', 0)

    # 姓名部分
    # http://python-docx.readthedocs.io/en/latest/dev/analysis/features/text/paragraph-format.html
    name = document.add_paragraph()
    # http://python-docx.readthedocs.io/en/latest/api/enum/WdAlignParagraph.html#wdparagraphalignment
    name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = name.add_run('日期：___________   姓名：___________')
    run.font.size = Pt(12)

    # 题目分类
    category_list = []
    for cg in QuestionBank.objects.all():
        if cg.category not in category_list:
            category_list.append(cg.category)
    # 提醒分类：选择、问答题
    questiontype_list = []
    for qt in QuestionBank.objects.all():
        if qt.questiontype not in questiontype_list:
            questiontype_list.append(qt.questiontype)

    if choose == '简单' or choose == '普通':
        if choose == '简单':
            word_name = 'IT工程师笔试题-简单.docx'
        else:
            word_name = 'IT工程师笔试题-普通.docx'

        for category in category_list:
            document.add_heading(category, 1)

            for questiontype in questiontype_list:
                count = 2
                if '问答' in questiontype:
                    count = 4
                questions = QuestionBank.objects.filter(level=choose, category=category, questiontype=questiontype)
                if questions.count() >= count:
                    rdm = random.sample(range(questions.count()), 2)
                    i = 1
                    for question_num in rdm:
                        question = '问题%s:%s' % (i, questions[question_num].text)
                        i += 1

                        paragraph = document.add_paragraph()
                        run = paragraph.add_run(question)
                        run.font.name = '宋体'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                        # 如果有图片
                        if questions[question_num].image != '':
                            # print(questions[question_num].image.url)
                            image_url = 'http://127.0.0.1' + questions[question_num].image.url
                            print(image_url)

                            req = request.Request(image_url)
                            try:
                                urlopen = request.urlopen(req, timeout=5)
                            except urllib.error.URLError:
                                print('连接错误')
                                image_url = 'http://127.0.0.1:8899' + questions[question_num].image.url
                                print(image_url)

                                req = request.Request(image_url)
                                urlopen = request.urlopen(req, timeout=5)

                            if urlopen.getcode() == 200:
                                imagedata = urlopen.read()
                                print('有图片导出！')
                                image_io = BytesIO()
                                image_io.write(imagedata)
                                image_io.seek(0)
                                # 直接保存原图
                                document.add_picture(image_io)
                                # document.add_picture(image_io, width=Inches(6))
                                # paragraph = document.add_paragraph('图片站位')

                        # 如果是问答题
                        if questions[question_num].questiontype == '问答题':
                            paragraph = document.add_paragraph('\n\n\n')

                # elif questions.count() < count and questions.count() > 0:
                elif 0 < questions.count() < count:
                    i = 1
                    for question_num in range(questions.count()):
                        question = '问题%s:%s' % (i, questions[question_num].text)
                        i += 1

                        paragraph = document.add_paragraph()
                        run = paragraph.add_run(question)
                        run.font.name = '宋体'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                        # 如果有图片
                        if questions[question_num].image != '':
                            image_url = 'http://127.0.0.1' + questions[question_num].image.url
                            print(image_url)

                            req = request.Request(image_url)
                            try:
                                urlopen = request.urlopen(req, timeout=5)
                            except urllib.error.URLError:
                                print('连接错误')
                                image_url = 'http://127.0.0.1:8899' + questions[question_num].image.url
                                print(image_url)

                                req = request.Request(image_url)
                                urlopen = request.urlopen(req, timeout=5)

                            if urlopen.getcode() == 200:
                                imagedata = urlopen.read()
                                print('有图片导出！')
                                image_io = BytesIO()
                                image_io.write(imagedata)
                                image_io.seek(0)
                                # 直接保存原图
                                document.add_picture(image_io)
                                # document.add_picture(image_io, width=Inches(6))
                                # paragraph = document.add_paragraph('图片站位')

                        # 如果是问答题
                        if questions[question_num].questiontype == '问答题':
                            paragraph = document.add_paragraph('\n\n\n')
    else:
        if choose == '所有简单':
            word_name = 'IT工程师笔试题-所有简单.docx'
            select_level = '简单'
        else:
            word_name = 'IT工程师笔试题-所有普通.docx'
            select_level = '普通'
        for category in category_list:
            document.add_heading(category, 1)
            questions = QuestionBank.objects.filter(level=select_level, category=category).order_by('category')
            i = 1
            for question_num in range(questions.count()):
                question = '问题%s:%s' % (i, questions[question_num].text)
                i += 1

                paragraph = document.add_paragraph()
                run = paragraph.add_run(question)
                run.font.name = '宋体'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 如果有图片
                if questions[question_num].image != '':
                    image_url = 'http://127.0.0.1' + questions[question_num].image.url
                    print(image_url)

                    req = request.Request(image_url)
                    try:
                        urlopen = request.urlopen(req, timeout=5)
                    except urllib.error.URLError:
                        print('连接错误')
                        image_url = 'http://127.0.0.1:8899' + questions[question_num].image.url
                        print(image_url)

                        req = request.Request(image_url)
                        urlopen = request.urlopen(req, timeout=5)

                    if urlopen.getcode() == 200:
                        imagedata = urlopen.read()
                        print('有图片导出！')
                        image_io = BytesIO()
                        image_io.write(imagedata)
                        image_io.seek(0)
                        # 直接保存原图
                        document.add_picture(image_io)
                        # document.add_picture(image_io, width=Inches(6))
                        # paragraph = document.add_paragraph('图片站位')

                # 如果是问答题
                if questions[question_num].questiontype == '问答题':
                    paragraph = document.add_paragraph('\n\n\n')

    # 保存文件到相应路径
    if platform.system() == 'Windows':
        word_path = settings.MEDIA_ROOT + '\\static\\questionbank'
        if not os.path.exists(word_path):
            os.makedirs(word_path)
        word_file = word_path + '\\' + word_name
        # 保存文件
        document.save(word_file)

    if platform.system() == 'Linux':
        word_path = settings.MEDIA_ROOT + '/static/questionbank'
        if not os.path.exists(word_path):
            os.mkdirs(word_path)
        word_file = word_path + '/' + word_name
        # 保存文件
        document.save(word_file)


if __name__ == '__main__':
    exportword('简单')